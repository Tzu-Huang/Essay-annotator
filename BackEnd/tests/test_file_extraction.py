import io
import unittest

from docx import Document
from pypdf import PdfWriter

from service.file_extraction import NoTextExtracted, UnsupportedFileType, extract_text


class FileExtractionTests(unittest.TestCase):
    def test_extracts_txt(self):
        text = extract_text("essay.txt", "Hello essay body.".encode("utf-8"))
        self.assertEqual(text, "Hello essay body.")

    def test_extracts_txt_with_bom(self):
        text = extract_text("essay.txt", "Hello essay body.".encode("utf-8-sig"))
        self.assertEqual(text, "Hello essay body.")

    def test_extracts_docx(self):
        buffer = io.BytesIO()
        document = Document()
        document.add_paragraph("First paragraph.")
        document.add_paragraph("Second paragraph.")
        document.save(buffer)

        text = extract_text("essay.docx", buffer.getvalue())
        self.assertIn("First paragraph.", text)
        self.assertIn("Second paragraph.", text)

    def test_docx_with_no_text_raises_no_text_extracted(self):
        buffer = io.BytesIO()
        Document().save(buffer)  # blank document, no paragraphs

        with self.assertRaises(NoTextExtracted):
            extract_text("empty.docx", buffer.getvalue())

    def test_pdf_with_no_text_layer_raises_no_text_extracted(self):
        # A PdfWriter with a blank page has no text layer -- this is the
        # "scanned image" case the design calls out explicitly.
        buffer = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.write(buffer)

        with self.assertRaises(NoTextExtracted):
            extract_text("scan.pdf", buffer.getvalue())

    def test_unsupported_extension_raises_unsupported_file_type(self):
        with self.assertRaises(UnsupportedFileType):
            extract_text("photo.png", b"not really a png")

    def test_legacy_doc_extension_is_unsupported(self):
        # Explicit non-goal from the design: only modern .docx, not legacy .doc.
        with self.assertRaises(UnsupportedFileType):
            extract_text("essay.doc", b"whatever bytes")

    def test_no_extension_raises_unsupported_file_type(self):
        with self.assertRaises(UnsupportedFileType):
            extract_text("essay", b"some text")

    def test_corrupt_docx_bytes_raise_unsupported_or_no_text(self):
        # Garbage bytes with a .docx extension must not raise an unhandled
        # exception type -- the caller (the endpoint) only catches these two.
        with self.assertRaises((UnsupportedFileType, NoTextExtracted)):
            extract_text("essay.docx", b"not a real docx file")


if __name__ == "__main__":
    unittest.main()
