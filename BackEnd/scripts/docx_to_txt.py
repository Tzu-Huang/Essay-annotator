# This is a general program that converts docx files to txt files

<<<<<<< HEAD
from pathlib import Path

from docx import Document


BACKEND_DIR = Path(__file__).resolve().parents[1]
INPUT_FILE = BACKEND_DIR / "drive_data" / "new_input"
OUTPUT_FILE = BACKEND_DIR / "drive_data" / "organized_data" / "new_input"


def docx_to_txt(docx_path: Path, txt_path: Path) -> None:
    doc = Document(docx_path)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    txt_path.write_text("\n".join(full_text), encoding="utf-8")


def main() -> None:
    OUTPUT_FILE.mkdir(parents=True, exist_ok=True)

    for docx_path in INPUT_FILE.glob("*.docx"):
        txt_path = OUTPUT_FILE / f"{docx_path.stem}.txt"
        docx_to_txt(docx_path, txt_path)
        print(f"Converted: {docx_path.name} -> {txt_path.name}")


if __name__ == "__main__":
    main()
=======
from docx import Document
import os

Input_file = 'Backend/drive_data/new_input/'
Output_file = 'Backend/drive_data/organized_data/new_input/'


def docx_to_txt(docx_path, txt_path):
    # load the Word document
    doc = Document(docx_path)
    
    # extract all paragraphs as plain text
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # join paragraphs with newlines and write to .txt
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))

def main():
    os.makedirs(Output_file, exist_ok=True)

    # for each file in the input directory
    for filename in os.listdir(Input_file):
        if filename.endswith('.docx'):
            
            docx_path = os.path.join(Input_file, filename)
            
            txt_filename = filename.replace('.docx', '.txt')
            txt_path = os.path.join(Output_file, txt_filename)
            
            docx_to_txt(docx_path, txt_path)
            print(f"Converted: {filename} → {txt_filename}")

if __name__ == "__main__":
    main()
>>>>>>> feature/admin
